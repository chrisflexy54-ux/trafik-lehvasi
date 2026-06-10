from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from PIL import Image

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Ara_Sinav_Donemi_Rapor_Trafik_Levhasi_YOLO26.docx"
OUTPUTS = ROOT / "outputs"
CAPTURES = ROOT / "captures"
ASSETS = OUTPUTS / "_report_assets"
ASSETS.mkdir(parents=True, exist_ok=True)

ACCENT = RGBColor(0xC0, 0x39, 0x2B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

def compress(src: Path, max_w: int = 1500, quality: int = 85) -> Path:
    im = Image.open(src).convert("RGB")
    if im.width > max_w:
        h = int(im.height * max_w / im.width)
        im = im.resize((max_w, h), Image.LANCZOS)
    dest = ASSETS / (src.stem + ".jpg")
    im.save(dest, "JPEG", quality=quality, optimize=True)
    return dest

def shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color
    })
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        run.font.color.rgb = color

def add_section_title(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. ")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = ACCENT
    run2 = p.add_run(text)
    run2.bold = True
    run2.font.size = Pt(15)
    run2.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "8",
        qn("w:space"): "2", qn("w:color"): "C0392B",
    })
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_stage_title(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    run2 = p.add_run(text)
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.color.rgb = DARK

def add_body(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_figure(doc, image_path: Path, width_in: float, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(compress(image_path)), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)

def add_placeholder(doc, caption: str):

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.2)
    shade(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("📷  EKRAN GÖRÜNTÜSÜ BURAYA EKLENECEK")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption)
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY

    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        b = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "dashed", qn("w:sz"): "10", qn("w:color"): "C0392B",
        })
        borders.append(b)
    tcPr.append(borders)
    doc.add_paragraph()

def add_caption_only(doc, caption: str):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)

def stage_visual(doc, user_capture: str | None, fallback: Path | None,
                 width_in: float, caption: str):

    if user_capture:
        path = CAPTURES / user_capture
        if path.exists():
            add_figure(doc, path, width_in, caption)
            return
    if fallback and fallback.exists():
        add_figure(doc, fallback, width_in, caption)
        return
    add_placeholder(doc, caption)

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=10)
        shade(hdr[i], "C0392B")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
            if i == 0:
                cells[i].paragraphs[0].runs[0].bold = True

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("NESNE TESPİTİ  ·  BİLGİSAYARLI GÖRÜ  ·  GERÇEK ZAMANLI")
    r.font.size = Pt(9); r.bold = True; r.font.color.rgb = ACCENT

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("ARA SINAV DÖNEMİ İLERLEME RAPORU")
    rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = DARK

    title2 = doc.add_paragraph(); title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt2 = title2.add_run("Trafik Levhası Tespiti — YOLO26 + SAHI")
    rt2.font.size = Pt(13); rt2.font.color.rgb = GREY

    make_table(doc, ["Bilgi", "Detay"], [
        ["Ad ve Soyad", "NTEKO CHRIS EMERY NGUESSAN"],
        ["Öğrenci Numarası", "223908708"],
        ["Proje", "Trafik Levhası Tespiti (YOLO26 + SAHI)"],
        ["Yaklaşım", "Klasik Tespit — Sınırlayıcı Kutu (Bounding Box)"],
        ["Son Gönderim Tarihi", "7 Haziran 2026"],
    ])
    doc.add_paragraph()

    add_section_title(doc, 1, "Projenin Amacı")
    add_body(doc,
        "Bu projenin amacı, yol kenarına yerleştirilmiş bir gözetleme kamerasından "
        "gelen video akışını analiz ederek her türlü trafik levhasını otomatik ve "
        "gerçek zamanlı biçimde tespit edebilen bir model tasarlamak ve eğitmektir. "
        "Sistem; hız sınırı, öncelik, yasak, tehlike ve yön levhaları ile trafik "
        "ışıklarını sahnede sürekli olarak tanıyabilmektedir.")

    add_section_title(doc, 2, "Kullanılan Teknolojiler")
    make_table(doc, ["Teknoloji", "Görevi"], [
        ["YOLO26 (Ultralytics)",
         "Gerçek zamanlı tespit; NMS'siz uçtan uca çıkarım, MuSGD, ProgLoss + STAL"],
        ["SAHI", "Dilimleyici Yardımlı Hiper Çıkarım — küçük/uzak levhalar"],
        ["PyTorch", "Derin öğrenme arka ucu"],
        ["OpenCV", "Video yakalama, annotasyon, kayıt"],
        ["Roboflow Universe", "YOLO formatında çok koşullu, annoteli veri setleri"],
        ["Python 3", "Boru hattının geliştirme dili"],
    ])
    doc.add_paragraph()

    add_section_title(doc, 3, "Benzer Projelere Göre Avantajlar")
    add_bullet(doc, "YOLO26: YOLO11'e kıyasla CPU çıkarımında %43'e varan hız kazancı.")
    add_bullet(doc, "NMS'siz uçtan uca çıkarım: daha düşük gecikme, gerçek zamana uygun.")
    add_bullet(doc, "SAHI: klasik tespitin kaçırdığı küçük/uzak levhaları yakalar.")
    add_bullet(doc, "Çok koşullu veri (gündüz, gece, yağmur, kar, bulanıklık) ile sağlamlık.")

    add_section_title(doc, 4, "Yazılım Geliştirme Aşamaları ve Ekran Görüntüleri")
    add_body(doc,
        "Aşağıda projenin her bir geliştirme aşaması, ilgili ekran görüntüsü ve "
        "yapılan çalışmanın açıklamasıyla birlikte sunulmaktadır.")

    add_stage_title(doc, "Aşama 1", "Proje Yapısı ve Kod Mimarisi")
    add_body(doc,
        "Proje, modüler bir yapı üzerine kurulmuştur: veri hazırlama, eğitim, "
        "değerlendirme, klasik çıkarım, SAHI çıkarımı ve gerçek zamanlı kamera için "
        "ayrı betikler bulunmaktadır. config.py dosyası, model adı bulunamazsa "
        "YOLO26'dan YOLO11'e otomatik geçiş (fallback) sağlamaktadır.")
    stage_visual(doc, "01_yapi.png", None, 6.0,
        "Şekil 1 — Editörde proje klasör yapısı (src/, configs/, outputs/, run_demo.py).")

    add_stage_title(doc, "Aşama 2", "Model Mimarisi (config.py / model_utils.py)")
    add_body(doc,
        "Bu aşamada modelin yapılandırması yazılmıştır: tercih sırası "
        "YOLO26 → YOLO11 → YOLOv8, güven ve IoU eşikleri, SAHI dilim boyutları. "
        "Böylece kod, hangi YOLO sürümü mevcutsa onunla sorunsuz çalışır.")
    stage_visual(doc, "02_kod_config.png", None, 6.0,
        "Şekil 2 — config.py / model_utils.py: model seçimi ve otomatik fallback kodu.")

    add_stage_title(doc, "Aşama 3", "Model Yükleme ve Uçtan Uca Demonun Çalıştırılması")
    add_body(doc,
        "Boru hattı ilk çalıştırıldığında YOLO26 ağırlıkları (yolo26n.pt) otomatik "
        "olarak indirilmiş ve yerel makinede başarıyla yüklenmiştir. run_demo.py "
        "betiği tüm boru hattını uçtan uca çalıştırır: model yükleme, örnek "
        "görüntüler, klasik ve SAHI tespiti. Terminal çıktısı hem modelin "
        "yüklendiğini hem de her sahne için klasik ve SAHI nesne sayılarını "
        "(örn. 3 → 12, 9 → 28) göstererek en güncel YOLO sürümünün fiilen "
        "çalıştığını doğrular.")
    stage_visual(doc, "03_terminal_demo.png", None, 6.0,
        "Şekil 3 — Terminal: yolo26n.pt indirilmesi, 'Modèle chargé : yolo26n.pt' "
        "ve run_demo.py çıktısı (Klasik/SAHI nesne sayıları).")

    add_stage_title(doc, "Aşama 4", "Klasik Tespit (YOLO26)")
    add_body(doc,
        "Statik bir görüntü üzerinde klasik YOLO26 çıkarımı uygulanmıştır. Model, "
        "araçları ve yakın nesneleri tespit etmekte; ancak uzaktaki küçük levhalar "
        "klasik yöntemde gözden kaçabilmektedir.")
    stage_visual(doc, None, OUTPUTS / "classic_stop_rue1.jpg", 5.5,
        "Şekil 4 — Klasik YOLO26 tespiti: yakın nesneler bulunur, uzaktaki DUR "
        "levhaları kaçırılır (3 nesne).")

    add_stage_title(doc, "Aşama 5", "SAHI ile Tespit (Dilimleme)")
    add_body(doc,
        "SAHI, görüntüyü çakışan dilimlere bölüp her birinde YOLO26 çalıştırarak "
        "sonuçları birleştirir. Bu sayede küçük ve uzak nesneler de yakalanır. "
        "Kentsel sahnede SAHI 28 nesne tespit etmiştir (trafik ışıkları + araçlar).")
    stage_visual(doc, None, OUTPUTS / "sahi_feux_ville.png", 3.6,
        "Şekil 5 — SAHI ile çoklu tespit (feux_ville): trafik ışıkları ve araçlar, "
        "28 nesne.")

    add_stage_title(doc, "Aşama 6", "Klasik vs SAHI Karşılaştırması")
    add_body(doc,
        "Aynı sahnede iki yöntem karşılaştırılmıştır. Sol tarafta klasik tespit "
        "DUR levhalarını kaçırırken, sağ tarafta SAHI bunları yakalamaktadır "
        "(3 nesne → 12 nesne). Bu, SAHI'nin somut katkısını göstermektedir.")
    stage_visual(doc, None, OUTPUTS / "comparaison_stop_rue1.jpg", 6.5,
        "Şekil 6 — Klasik YOLO26 (sol) vs YOLO26 + SAHI (sağ): SAHI uzaktaki DUR "
        "levhalarını yakalar (3 → 12 nesne).")

    add_stage_title(doc, "Aşama 7", "Eğitim ve Değerlendirme Betikleri")
    add_body(doc,
        "train.py, YOLO26'yı 'auto' optimizör (MuSGD) ve erken durdurma ile eğitir; "
        "evaluate.py ise mAP50, mAP50-95, hassasiyet ve geri çağırma metriklerini "
        "hesaplar. Betikler hazırdır; nihai eğitim, veri seti birleştirildikten "
        "sonra GPU üzerinde yapılacaktır.")
    stage_visual(doc, "05_kod_train.png", None, 5.5,
        "Şekil 7 — train.py: YOLO26 eğitim betiği (epochs, MuSGD, erken durdurma).")

    add_stage_title(doc, "Aşama 8", "Gerçek Zamanlı Kamera Tespiti")
    add_body(doc,
        "realtime_camera.py, webcam akışını kare kare analiz ederek levhaları canlı "
        "olarak işaretler ve FPS gösterir — bir trafik gözetleme kamerasının "
        "koşullarını simüle eder. Aynı tespit motoru, aşağıdaki yol görüntüsünde "
        "olduğu gibi her karede çalışır.")
    stage_visual(doc, None, OUTPUTS / "sahi_stop_rue2.png", 5.4,
        "Şekil 8 — Gerçek yol sahnesinde tespit (gerçek zamanlı mod ile aynı motor): "
        "uzaktaki DUR levhası yakalanmıştır.")

    add_section_title(doc, 5, "Mevcut Durum Özeti")
    make_table(doc, ["Aşama", "Durum"], [
        ["Boru hattı mimarisi ve teknoloji seçimi", "Tamamlandı"],
        ["Klasik + SAHI çıkarım betikleri", "Tamamlandı (çalışıyor)"],
        ["Gerçek zamanlı kamera tespiti", "Tamamlandı (betik hazır)"],
        ["Uçtan uca demo ve görsel üretimi", "Tamamlandı"],
        ["Klasik vs SAHI karşılaştırması", "Doğrulandı (3→12, 9→28)"],
        ["Eğitim + değerlendirme betikleri", "Kodlandı, çalıştırmaya hazır"],
        ["Roboflow çok koşullu veri seti birleştirme", "Sıradaki adım"],
        ["GPU üzerinde nihai eğitim + metrikler", "Sıradaki adım"],
    ])
    doc.add_paragraph()

    add_section_title(doc, 6, "Sonuç")
    add_body(doc,
        "Ara sınav dönemine kadar projenin yazılım iskeleti tamamlanmış ve "
        "YOLO26 + SAHI yaklaşımı gerçek görüntüler üzerinde çalışır biçimde "
        "doğrulanmıştır. Sonraki aşama, Roboflow Universe'den çok koşullu veri "
        "setinin birleştirilmesi ve GPU üzerinde nihai modelin eğitilmesidir.")

    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = foot.add_run("YOLO26 Projesi  ·  Trafik Levhası Tespiti  ·  "
                      "NTEKO CHRIS EMERY NGUESSAN  ·  223908708")
    rf.font.size = Pt(8); rf.font.color.rgb = GREY

    doc.save(OUT)
    print(f"[rapor] Oluşturuldu: {OUT}")

if __name__ == "__main__":
    build()
